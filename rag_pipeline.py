"""
RAG Pipeline - Person 1's Component
Handles document processing, embedding creation, and retrieval
"""

import os
import pickle
from dataclasses import dataclass
from typing import List, Dict, Any
import PyPDF2
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


@dataclass
class Document:
    """Minimal document structure used by the app."""
    page_content: str
    metadata: Dict[str, Any]


class RAGPipeline:
    """
    Retrieval-Augmented Generation Pipeline
    Processes documents, creates embeddings, and retrieves relevant content
    """
    
    def __init__(self):
        """
        Initialize RAG pipeline with TF-IDF vectorization.
        """
        self.vectorizer = TfidfVectorizer(stop_words="english", ngram_range=(1, 2))
        self.doc_matrix = None
        self.documents = []
        self.embedding_dim = 0
        
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text() or ""
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"Error extracting PDF with PyPDF2: {e}")
        
        return text.strip()
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(docx_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text.strip()
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
            return ""
    
    def extract_text_from_txt(self, txt_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            print(f"Error reading TXT file: {e}")
            return ""

    @staticmethod
    def split_text(text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> List[str]:
        """Split text into overlapping chunks."""
        if not text:
            return []
        if chunk_overlap >= chunk_size:
            chunk_overlap = 0
        chunks = []
        start = 0
        step = chunk_size - chunk_overlap
        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            start += step
        return chunks
    
    def process_document(self, file_path: str) -> List[Document]:
        """
        Process uploaded document and split into chunks
        
        Args:
            file_path: Path to the document
            
        Returns:
            List of Document objects with text chunks
        """
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Extract text based on file type
        if file_extension == '.pdf':
            text = self.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            text = self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            text = self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        if not text:
            raise ValueError("No text could be extracted from the document")
        
        # Split text into chunks
        chunks = self.split_text(text)
        
        # Create Document objects
        documents = [
            Document(page_content=chunk, metadata={"source": file_path, "chunk_id": i})
            for i, chunk in enumerate(chunks)
        ]
        
        return documents
    
    def build_index(self, documents: List[Document]):
        """
        Build TF-IDF index from documents.
        
        Args:
            documents: List of Document objects
        """
        if not documents:
            return

        self.documents.extend(documents)
        texts = [doc.page_content for doc in self.documents]
        self.doc_matrix = self.vectorizer.fit_transform(texts)
        self.embedding_dim = self.doc_matrix.shape[1]
        
    def retrieve(self, query: str, top_k: int = 3) -> List[Document]:
        """
        Retrieve most relevant documents for a query
        
        Args:
            query: Search query
            top_k: Number of documents to retrieve
            
        Returns:
            List of most relevant Document objects
        """
        if self.doc_matrix is None or len(self.documents) == 0:
            return []

        query_vector = self.vectorizer.transform([query])
        scores = cosine_similarity(query_vector, self.doc_matrix).flatten()
        ranked_indices = scores.argsort()[::-1]

        relevant_docs = []
        for idx in ranked_indices:
            if scores[idx] <= 0:
                continue
            relevant_docs.append(self.documents[idx])
            if len(relevant_docs) >= top_k:
                break
        
        return relevant_docs
    
    def save_vectorstore(self, save_path: str = "embeddings/vectorstore.pkl"):
        """Save the vector store to disk"""
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        
        data = {
            'vectorizer': self.vectorizer,
            'doc_matrix': self.doc_matrix,
            'documents': self.documents
        }
        
        with open(save_path, 'wb') as f:
            pickle.dump(data, f)
        
        print(f"Vector store saved to {save_path}")
    
    def load_vectorstore(self, load_path: str = "embeddings/vectorstore.pkl"):
        """Load the vector store from disk"""
        if not os.path.exists(load_path):
            print(f"No vector store found at {load_path}")
            return False
        
        with open(load_path, 'rb') as f:
            data = pickle.load(f)
        
        self.vectorizer = data.get('vectorizer', TfidfVectorizer(stop_words="english", ngram_range=(1, 2)))
        self.doc_matrix = data.get('doc_matrix')
        self.documents = data.get('documents', [])
        if self.doc_matrix is not None:
            self.embedding_dim = self.doc_matrix.shape[1]
        
        print(f"Vector store loaded from {load_path}")
        return True
    
    def get_context_for_query(self, query: str, top_k: int = 3) -> str:
        """
        Get concatenated context from retrieved documents
        
        Args:
            query: User query
            top_k: Number of documents to retrieve
            
        Returns:
            Concatenated text from relevant documents
        """
        relevant_docs = self.retrieve(query, top_k)
        context = "\n\n".join([doc.page_content for doc in relevant_docs])
        return context


# Test function
if __name__ == "__main__":
    # Example usage
    rag = RAGPipeline()
    
    # Test with a sample document
    print("RAG Pipeline initialized successfully!")
    print(f"Embedding dimension: {rag.embedding_dim}")
